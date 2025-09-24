import pdfplumber
import docx
import os

# ---------- PDF Extraction ----------
def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file with error handling."""
    text = ''
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    print(f"[WARN] Page {i+1} in {os.path.basename(pdf_path)} had no extractable text")
    except Exception as e:
        print(f"[ERROR] Failed extracting PDF {pdf_path}: {e}")
        return ""
    
    return text.strip()

# ---------- DOCX Extraction ----------
def extract_text_from_docx(docx_path):
    """Extract text from DOCX file with error handling."""
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        # Extract text from tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text

        return text.strip()
    except Exception as e:
        print(f"[ERROR] Failed extracting DOCX {docx_path}: {e}")
        return ""

# ---------- Dispatcher ----------
def extract_resume_text(file_path, filename):
    """Main function to extract text from resumes (PDF or DOCX)."""
    if filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif filename.lower().endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        print(f"[ERROR] Unsupported file format: {filename}")
        return ""

    if not validate_resume_text(text):
        print(f"[WARN] Resume {filename} may not contain meaningful text")
    return text

# ---------- Validation ----------
def validate_resume_text(text):
    """Validate if extracted text is meaningful."""
    if not text or len(text) < 100:
        return False

    # Look for resume-related keywords
    resume_keywords = [
        'experience', 'education', 'skills',
        'projects', 'work', 'university',
        'college', 'degree'
    ]
    text_lower = text.lower()
    keyword_count = sum(1 for kw in resume_keywords if kw in text_lower)

    return keyword_count >= 2  # Require at least 2 relevant keywords
